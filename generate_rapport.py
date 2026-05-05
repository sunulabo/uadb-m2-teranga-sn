# generate_rapport.py
# Generation du rapport final Teranga-SN
# NDIAYE Papa Malick - Eq.12 - Master 2 DSGL UADB 2025-2026

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')
OUT_PATH = os.path.join(BASE_DIR, 'rapport_teranga_sn.docx')


def para(doc, texte, taille=11, gras=False, italique=False, centre=False):
    p = doc.add_paragraph()
    if centre:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texte)
    run.font.size = Pt(taille)
    run.bold = gras
    run.italic = italique
    return p


def puce(doc, texte):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(texte)
    run.font.size = Pt(11)
    return p


def titre1(doc, texte):
    h = doc.add_heading(texte, level=1)
    h.runs[0].font.size = Pt(14)
    h.runs[0].bold = True
    return h


def titre2(doc, texte):
    h = doc.add_heading(texte, level=2)
    h.runs[0].font.size = Pt(12)
    h.runs[0].bold = True
    return h


def espace(doc):
    doc.add_paragraph('')


def tableau(doc, entetes, lignes, largeurs=None):
    t = doc.add_table(rows=1 + len(lignes), cols=len(entetes))
    t.style = 'Table Grid'
    # En-tetes en gras, fond blanc
    for i, h in enumerate(entetes):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    # Donnees
    for r, ligne in enumerate(lignes):
        for c, val in enumerate(ligne):
            cell = t.rows[r + 1].cells[c]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    # Largeurs
    if largeurs:
        for row in t.rows:
            for i, w in enumerate(largeurs):
                row.cells[i].width = Cm(w)
    return t


def bloc_code(doc, texte):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p


def page_titre(doc):
    espace(doc)
    para(doc, 'UNIVERSITE ALIOUNE DIOP DE BAMBEY', taille=13, gras=True, centre=True)
    para(doc, 'Ecole Superieure de Technologie', taille=11, centre=True)
    para(doc, 'Master 2 - Sciences et Gestion du Logiciel et des Donnees (DSGL)', taille=11, centre=True)
    espace(doc)
    espace(doc)
    para(doc, 'PROJET FINAL - BIG DATA', taille=13, gras=True, centre=True)
    espace(doc)
    para(doc, 'TERANGA-SN', taille=22, gras=True, centre=True)
    para(doc, 'Plateforme d\'Intelligence Touristique et d\'Economie Numerique au Senegal', taille=13, centre=True)
    espace(doc)
    espace(doc)
    tableau(doc,
        ['', ''],
        [
            ['Equipe', 'Eq.12'],
            ['Membres', 'NDIAYE Papa Malick  |  TINE Abdoussalam'],
            ['Encadrant', 'Professeur Big Data - UADB'],
            ['Annee universitaire', '2025-2026'],
            ['Date de rendu', 'Mai 2026'],
        ],
        largeurs=[5, 11]
    )
    espace(doc)
    img = os.path.join(DATA_DIR, 'dashboard_teranga_sn.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(5.8))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para(doc, 'Apercu du tableau de bord Teranga-SN (mode demonstration)', taille=9, italique=True, centre=True)
    doc.add_page_break()


def resume(doc):
    titre1(doc, 'Resume executif')
    para(doc,
        'Teranga-SN est une plateforme Big Data deployee pour la Direction Generale du '
        'Tourisme et des Transports Touristiques (DGTT) du Senegal. Elle collecte, '
        'traite et analyse en temps reel deux flux de donnees : les avis touristiques '
        'publies en francais, anglais et wolof, et les transactions e-commerce des '
        'plateformes locales (Jumia, Expat-Dakar, Senmarket).')
    espace(doc)
    para(doc,
        'Le systeme repose sur une architecture Big Data complete : Apache Kafka assure '
        'l\'ingestion, Apache Spark Structured Streaming traite les donnees en flux, '
        'HBase et Hive assurent le stockage, un modele Random Forest predit les flux '
        'de touristes, et Apache Airflow orchestre le tout de facon hebdomadaire.')
    espace(doc)
    para(doc,
        'Les resultats techniques sont les suivants : 10 tests unitaires valides (10/10), '
        'un modele de prediction avec un R2 de 0.90, et un pipeline complet de 8 '
        'services Docker operationnels. Ce projet repond a un besoin reel et strategique '
        'pour le secteur touristique senegalais.')
    doc.add_page_break()


def introduction(doc):
    titre1(doc, '1. Introduction et Contexte')

    para(doc,
        'Le tourisme represente environ 7% du PIB du Senegal et constitue l\'une des '
        'premieres sources de devises du pays. Pourtant, les acteurs institutionnels '
        'comme la DGTT disposent de peu d\'outils modernes pour analyser les flux de '
        'visiteurs en temps reel ou anticiper les periodes de crise de reputation.')
    espace(doc)
    para(doc,
        'Le projet Teranga-SN repond directement a ce manque. Son nom vient du mot '
        'wolof qui signifie hospitalite, valeur fondamentale de la culture senegalaise '
        'et premier attrait pour les touristes etrangers. L\'idee est simple : si nous '
        'pouvons mesurer la satisfaction en temps reel sur toutes les destinations du '
        'pays et dans toutes les langues parles par les visiteurs, nous pouvons agir '
        'avant qu\'une mauvaise reputation ne fasse fuir les touristes.')
    espace(doc)

    titre2(doc, '1.1 Problematiques traitees')
    para(doc, 'Ce projet repond a trois problematiques concretes :')
    puce(doc, 'Comment analyser des avis touristiques en wolof alors que la plupart des outils NLP ne couvrent pas cette langue ?')
    puce(doc, 'Comment proteger les donnees personnelles des touristes tout en conservant la possibilite de detecter les doublons et le spam ?')
    puce(doc, 'Comment prevoir le nombre d\'arrivees touristiques par destination pour aider la DGTT a planifier ses ressources ?')
    espace(doc)

    titre2(doc, '1.2 Objectifs du projet')
    tableau(doc,
        ['Objectif', 'Description', 'Technologie utilisee'],
        [
            ['Ingestion temps reel', 'Collecter les avis et transactions en continu', 'Apache Kafka + NiFi'],
            ['Analyse NLP multilingue', 'Scorer le sentiment en FR, EN et WO', 'Spark Streaming + Lexique metier'],
            ['Privacy by Design', 'Anonymiser les donnees personnelles des la collecte', 'SHA-256 + sel cryptographique'],
            ['Stockage hybride', 'NoSQL pour le temps reel, SQL pour l\'analytique', 'HBase + Hive ORC/SNAPPY'],
            ['Prediction ML', 'Prevoir les flux touristiques par destination', 'Random Forest + MLflow'],
            ['Orchestration', 'Declencher le reentainement si la reputation chute', 'Airflow DAG hebdomadaire'],
            ['Visualisation', 'Tableau de bord pour les decideurs de la DGTT', 'Matplotlib + Seaborn'],
        ],
        largeurs=[4, 7, 5]
    )
    doc.add_page_break()


def architecture(doc):
    titre1(doc, '2. Architecture Technique')

    para(doc,
        'L\'architecture de Teranga-SN suit un pattern Lambda simplifie. Les donnees '
        'entrent par Kafka, sont traitees par Spark en temps reel, puis stockees dans '
        'deux systemes complementaires : HBase pour les requetes rapides operationnelles, '
        'et Hive pour les analyses historiques. Airflow surveille le tout chaque semaine.')
    espace(doc)

    titre2(doc, '2.1 Flux de donnees de bout en bout')
    para(doc,
        'Voici comment une donnee voyage dans le systeme, depuis sa creation jusqu\'a '
        'son affichage dans le tableau de bord :')
    espace(doc)
    para(doc,
        '1. Un touriste publie un avis sur TripAdvisor. NiFi ou le producteur Kafka '
        'capture cet avis et l\'envoie dans le topic teranga_avis_raw.')
    para(doc,
        '2. Spark Structured Streaming lit le topic. Il anonymise immediatement le '
        'user_id avec SHA-256, analyse le sentiment du texte, et calcule un score.')
    para(doc,
        '3. Les donnees enrichies et anonymisees sont ecrites dans HBase pour '
        'consultation rapide et dans Hive pour l\'historique analytique.')
    para(doc,
        '4. Chaque lundi, le DAG Airflow interroge Hive. Si plus d\'une destination '
        'est en alerte ROUGE, il relance automatiquement l\'entrainement du modele ML.')
    para(doc,
        '5. Le tableau de bord lit Hive et genere les quatre panneaux de visualisation '
        'pour les decideurs de la DGTT.')
    espace(doc)

    titre2(doc, '2.2 Infrastructure Docker')
    para(doc,
        'L\'ensemble des services est deploye via Docker Compose sur un reseau prive '
        'teranga-net. Le tableau suivant presente l\'etat de chaque service apres '
        'demarrage et les tests de validation realises :')
    espace(doc)
    tableau(doc,
        ['Service', 'Image Docker', 'Port', 'Etat valide'],
        [
            ['Zookeeper', 'confluentinc/cp-zookeeper:7.4.0', '2181', 'Healthy'],
            ['Kafka', 'confluentinc/cp-kafka:7.4.0', '9092 / 9093', 'Healthy - 2 topics crees'],
            ['NiFi', 'apache/nifi:1.23.2', '8081', 'Running - interface accessible'],
            ['HBase', 'harisekhon/hbase:2.1', '9090 / 16010', 'Healthy - 3 tables creees'],
            ['Hive', 'apache/hive:3.1.3', '10000', 'Running - 2 tables + 2 vues'],
            ['Spark Master', 'apache/spark:latest (v4.1.1)', '8080 / 7077', 'Running - worker enregistre'],
            ['Spark Worker', 'apache/spark:latest (v4.1.1)', 'interne', 'Running - 2 Go RAM, 2 coeurs'],
            ['Airflow', 'apache/airflow:2.7.3', '8082', 'Running - DAG charge'],
        ],
        largeurs=[3.5, 5.5, 3, 4]
    )
    doc.add_page_break()


def kafka_nifi(doc):
    titre1(doc, '3. Ingestion des Donnees : Kafka et NiFi')

    para(doc,
        'La couche d\'ingestion est le point d\'entree de toutes les donnees dans '
        'le systeme. Elle doit etre capable de recevoir des milliers de messages par '
        'jour de facon fiable et sans perte.')
    espace(doc)

    titre2(doc, '3.1 Architecture Kafka')
    para(doc,
        'Apache Kafka joue le role de bus de messages central. Deux topics ont ete '
        'crees, chacun correspondant a un type de donnee distinct :')
    espace(doc)
    tableau(doc,
        ['Topic Kafka', 'Source', 'Donnees transportees', 'Volume'],
        [
            ['teranga_avis_raw', 'NiFi / Producteur Python', 'Avis touristiques en JSON (FR, EN, WO)', 'Environ 1 message toutes les 3 secondes'],
            ['teranga_arrivees_raw', 'Producteur Python', 'Transactions e-commerce + arrivees DGTT', 'Environ 2 messages toutes les 3 secondes'],
        ],
        largeurs=[4, 3.5, 5, 3.5]
    )
    espace(doc)
    para(doc,
        'Pourquoi deux topics separes ? Parce que les avis et les transactions ont des '
        'schemas differents et des traitements differents dans Spark. Separer les flux '
        'des l\'ingestion simplifie le code et permet de scaler chaque topic '
        'independamment si le volume augmente.')
    espace(doc)

    titre2(doc, '3.2 Structure d\'un message brut')
    para(doc,
        'Voici un exemple de message tel qu\'il arrive dans Kafka, avant tout traitement. '
        'On remarque la presence de donnees personnelles (user_id, email_client) qui '
        'seront supprimees des la premiere etape de Spark :')
    espace(doc)
    bloc_code(doc,
        'Topic: teranga_avis_raw\n\n'
        '{\n'
        '  "avis_id":      "b29b6f5d-5762-44b4-a449-3f5267bcbd11",\n'
        '  "user_id":      "USR_27f989caa4",          <- donnee personnelle a anonymiser\n'
        '  "email_client": "user1680@test.sn",         <- donnee personnelle a supprimer\n'
        '  "destination":  "TOUBA",\n'
        '  "note":         1.4,\n'
        '  "texte_avis":   "Arnaque taxi, prix non affiche a l\'avance",\n'
        '  "langue":       "FR",\n'
        '  "timestamp":    "2026-05-05T18:41:22Z"\n'
        '}'
    )
    espace(doc)

    titre2(doc, '3.3 Role de NiFi')
    para(doc,
        'Apache NiFi est l\'outil d\'integration qui permet de connecter des sources '
        'externes comme l\'API TripAdvisor directement a Kafka, sans ecrire de code. '
        'Il est accessible sur http://localhost:8081 avec les identifiants admin/teranga2025!. '
        'Le template NiFi de l\'equipe (template_nifi_eq12.xml) definit un flux a '
        'trois processeurs : GenerateFlowFile simule les avis, et deux ProcessorPublishKafka '
        'envoient vers chacun des deux topics.')
    doc.add_page_break()


def spark(doc):
    titre1(doc, '4. Traitement en Flux : Spark Structured Streaming')

    para(doc,
        'Spark Structured Streaming est le coeur du systeme. C\'est lui qui transforme '
        'les donnees brutes en informations exploitables. Il realise trois operations '
        'en parallele : l\'anonymisation des donnees personnelles, l\'analyse de sentiment, '
        'et le calcul de la reputation par destination.')
    espace(doc)

    titre2(doc, '4.1 Privacy by Design : l\'anonymisation SHA-256')
    para(doc,
        'Des leur arrivee dans Spark, les donnees personnelles sont traitees selon le '
        'principe du Privacy by Design : on ne conserve jamais ce qu\'on n\'a pas le '
        'droit de stocker.')
    espace(doc)
    para(doc,
        'Pourquoi SHA-256 plutot que de simplement supprimer le user_id ? Parce que '
        'supprimer l\'identifiant completement nous empeche de detecter si un meme '
        'utilisateur envoie dix avis negatifs en une heure pour nuire a une destination. '
        'SHA-256 avec un sel cryptographique nous donne un identifiant stable et '
        'irreversible : on peut compter les recurrences sans jamais connaitre l\'identite reelle.')
    espace(doc)
    bloc_code(doc,
        '# Anonymisation : user_id -> SHA-256(user_id + sel) -> user_secure (64 caracteres hex)\n'
        '.withColumn("user_secure",\n'
        '    sha2(concat(col("user_id"), lit(SALT)), 256))\n'
        '.drop("user_id", "email_client")   # suppression immediate et definitive'
    )
    espace(doc)

    titre2(doc, '4.2 Analyse de sentiment multilingue')
    para(doc,
        'L\'analyse de sentiment utilise une approche lexicale adaptee au domaine '
        'touristique senegalais. Un dictionnaire de termes positifs et negatifs a ete '
        'construit en trois langues. Pour chaque avis, le score est la moyenne des '
        'valeurs des termes trouves dans le texte.')
    espace(doc)
    tableau(doc,
        ['Langue', 'Termes positifs (exemples)', 'Termes negatifs (exemples)'],
        [
            ['Francais (FR)', 'magnifique (+0.90), teranga (+0.80), excellent (+0.90)', 'arnaque (-0.90), sale (-0.80), decevant (-0.70)'],
            ['Anglais (EN)', 'amazing (+0.90), paradise (+0.90), beautiful (+0.80)', 'disappointed (-0.70), dirty (-0.80), rude (-0.75)'],
            ['Wolof (WO)', 'dafa baax (+0.90), rafet (+0.80), baax (+0.85)', 'dafa neka (-0.80), amul solo (-0.70)'],
        ],
        largeurs=[3, 7.5, 5.5]
    )
    espace(doc)
    para(doc, 'Resultats de validation sur 5 avis de test :')
    espace(doc)
    tableau(doc,
        ['Langue', 'Extrait de l\'avis', 'Score calcule', 'Label attribue'],
        [
            ['FR', 'Magnifique sejour, teranga incroyable !', '+1.70', 'POSITIF'],
            ['EN', 'Too many touts, disappointed', '-1.20', 'NEGATIF'],
            ['WO', 'Dafa baax torop, rafet na', '+2.55', 'POSITIF'],
            ['FR', 'Plage sale, arnaque taxi', '-1.70', 'NEGATIF'],
            ['EN', 'Amazing beaches, would recommend!', '+1.65', 'POSITIF'],
        ],
        largeurs=[2, 7.5, 3, 3.5]
    )
    espace(doc)

    titre2(doc, '4.3 Agregation de la reputation par fenetre glissante')
    para(doc,
        'La reputation d\'une destination est calculee sur les 7 derniers jours, '
        'avec une mise a jour toutes les 24 heures. Ce choix de fenetre a ete fait '
        'deliberement : une fenetre trop courte (1 jour) serait trop sensible aux '
        'variations aleatoires, une fenetre trop longue (30 jours) masquerait une '
        'crise qui eclate en quelques jours.')
    espace(doc)
    para(doc,
        'Le watermark de 1 jour resout un probleme technique important : sans lui, '
        'Spark garde en memoire l\'etat de toutes les fenetres passees indefiniment. '
        'Sur un flux continu de 24h/24, cela provoquerait un OutOfMemoryError. '
        'Le watermark dit a Spark : au-dela d\'un jour de retard, ignore l\'evenement '
        'et libere la memoire de cette fenetre.')
    espace(doc)
    para(doc, 'Regles de classification de la reputation :')
    tableau(doc,
        ['Statut', 'Condition sur le score moyen', 'Action declenchee'],
        [
            ['VERT', 'Score moyen >= 0.0', 'Aucune alerte'],
            ['ORANGE', 'Score moyen entre -0.3 et 0.0', 'Alerte ecrite dans HBase'],
            ['ROUGE', 'Score moyen < -0.3', 'Alerte HBase + reentainement ML si 2+ destinations'],
        ],
        largeurs=[2.5, 6, 7.5]
    )
    doc.add_page_break()


def stockage(doc):
    titre1(doc, '5. Stockage des Donnees : HBase et Hive')

    para(doc,
        'Le projet utilise deux systemes de stockage complementaires. HBase pour '
        'les acces rapides en lecture/ecriture operationnelle, et Hive pour les '
        'requetes analytiques SQL sur l\'historique. Les deux ont des forces '
        'differentes et se completent naturellement.')
    espace(doc)

    titre2(doc, '5.1 HBase : le stockage NoSQL operationnel')
    para(doc,
        'HBase est une base de donnees NoSQL orientee colonnes qui peut ecrire et '
        'lire des millions de lignes par seconde. Dans Teranga-SN, il stocke les '
        'donnees que Spark produit en temps reel et les alertes de reputation.')
    espace(doc)
    para(doc, 'Trois tables ont ete creees dans le namespace "teranga" :')
    espace(doc)
    tableau(doc,
        ['Table HBase', 'Familles de colonnes', 'Contenu stocke'],
        [
            ['teranga:avis_analyses', 'meta, sentiment, geo', 'Chaque avis apres enrichissement Spark (destination, note, score, label)'],
            ['teranga:alertes_reputation', 'alerte', 'Destinations en statut ROUGE ou ORANGE avec horodatage'],
            ['teranga:stats_ecommerce', 'stats, meta', 'Agregats e-commerce par destination (CA, nb transactions)'],
        ],
        largeurs=[5, 4, 7]
    )
    espace(doc)
    para(doc,
        'Le test d\'ecriture et de lecture realise confirme le bon fonctionnement : '
        'une ligne inseree avec 9 colonnes est restituee integralement par la lecture '
        'suivante. La connexion utilise le protocole Thrift sur le port 9090.')
    espace(doc)

    titre2(doc, '5.2 Hive : l\'entrepot analytique SQL')
    para(doc,
        'Hive permet d\'interroger les donnees historiques avec du SQL standard. '
        'Les tables sont stockees au format ORC avec compression SNAPPY, ce qui '
        'divise la taille des fichiers par 5 par rapport au format texte brut '
        'et accelere les requetes analytiques.')
    espace(doc)
    para(doc, 'Le schema Hive de Teranga-SN comprend deux tables et deux vues :')
    espace(doc)
    tableau(doc,
        ['Objet', 'Type', 'Description et justification technique'],
        [
            ['avis_analyses', 'TABLE', 'Partitionnee par date_obs (YYYY-MM-DD). Le partitionnement accelere les requetes sur les 30 derniers jours car Hive ne lit que les partitions concernees.'],
            ['ecommerce_transactions', 'TABLE', 'Transactions e-commerce anonymisees, format ORC.'],
            ['vue_destinations', 'VUE SQL', 'Calcule automatiquement la reputation de chaque destination sur les 30 derniers jours avec son statut ROUGE/ORANGE/VERT.'],
            ['vue_ecommerce', 'VUE SQL', 'Classement des categories de produits par chiffre d\'affaires et par region de livraison.'],
        ],
        largeurs=[4, 2.5, 9.5]
    )
    doc.add_page_break()


def validation(doc):
    titre1(doc, '6. Validation et Qualite des Donnees')

    para(doc,
        'La qualite des donnees est un enjeu critique dans un systeme Big Data. '
        'Une donnee corrompue ou mal anonymisee peut compromettre l\'analyse entiere '
        'ou violer le RGPD. Nous avons mis en place deux niveaux de validation.')
    espace(doc)

    titre2(doc, '6.1 Schemas Pandera : validation automatique')
    para(doc,
        'Pandera permet de definir des schemas de validation declaratifs pour les '
        'DataFrames pandas. Chaque fois qu\'un batch de donnees entre dans la couche '
        'de stockage, il est valide contre le schema correspondant. Si une contrainte '
        'est violee, une exception est levee et les donnees ne sont pas ecrites.')
    espace(doc)
    para(doc,
        'Un apport original de notre implementation : la validation verifie non '
        'seulement les valeurs, mais aussi l\'absence de certaines colonnes. Si '
        'user_id ou email_client sont encore presents dans le DataFrame a cette '
        'etape, c\'est que l\'anonymisation a echoue. Le schema le detecte et bloque '
        'l\'ingestion.')
    espace(doc)

    titre2(doc, '6.2 Resultats de la suite de tests (10/10)')
    tableau(doc,
        ['Test', 'Cas teste', 'Comportement attendu', 'Resultat'],
        [
            ['test_avis_valide', 'DataFrame conforme', 'Validation OK, retour du DataFrame', 'PASSE'],
            ['test_avis_note_hors_intervalle', 'note = 6.0 (maximum = 5.0)', 'SchemaErrors leve', 'PASSE'],
            ['test_avis_destination_inconnue', 'destination = "PARIS"', 'SchemaErrors leve', 'PASSE'],
            ['test_avis_pii_user_id_detecte', 'colonne user_id presente', 'SchemaErrors leve', 'PASSE'],
            ['test_avis_pii_email_detecte', 'colonne email_client presente', 'SchemaErrors leve', 'PASSE'],
            ['test_avis_sentiment_label_invalide', 'label = "TRES_POSITIF"', 'SchemaErrors leve', 'PASSE'],
            ['test_ecomm_valide', 'Transaction conforme', 'Validation OK', 'PASSE'],
            ['test_ecomm_montant_negatif', 'montant = -500 FCFA', 'SchemaErrors leve', 'PASSE'],
            ['test_ecomm_categorie_inconnue', 'categorie = "VOITURE"', 'SchemaErrors leve', 'PASSE'],
            ['test_ecomm_pii_user_id_detecte', 'user_id present (e-commerce)', 'SchemaErrors leve', 'PASSE'],
        ],
        largeurs=[5, 4, 4.5, 2.5]
    )
    espace(doc)
    bloc_code(doc,
        '$ python -m pytest tests/ -v\n'
        '======== 10 passed, 1 warning in 1.56s ========'
    )
    espace(doc)

    titre2(doc, '6.3 Validation du pipeline complet')
    tableau(doc,
        ['Composant', 'Test realise', 'Resultat observe'],
        [
            ['Kafka producer', 'Envoi de messages FR, EN et WO pendant 12 secondes', '8 messages dans teranga_avis_raw, 10 dans teranga_arrivees_raw'],
            ['HBase ecriture', 'Insertion d\'un avis TOUBA avec 9 colonnes', 'Ligne inseree et relue integralement - OK'],
            ['HBase alertes', 'Insertion d\'une alerte ROUGE pour TOUBA', 'Alerte stockee avec horodatage - OK'],
            ['Hive schema', 'Chargement du fichier hive_setup.sql', 'Base teranga_sn, 2 tables, 2 vues creees - OK'],
            ['Spark master', 'Demarrage du cluster standalone', 'Spark 4.1.1, worker de 2 Go enregistre - OK'],
            ['Dashboard', 'Generation du PNG en mode demonstration', 'Fichier 299 Ko genere sans erreur - OK'],
        ],
        largeurs=[3.5, 6, 6.5]
    )
    doc.add_page_break()


def machine_learning(doc):
    titre1(doc, '7. Machine Learning : Prediction des Flux Touristiques')

    para(doc,
        'L\'objectif du modele de machine learning est de predire le nombre de '
        'touristes attendus par destination et par jour. Cette prediction permet '
        'a la DGTT d\'anticiper les periodes de forte affluence et d\'adapter '
        'ses ressources (guides, securite, capacite hoteliere).')
    espace(doc)

    titre2(doc, '7.1 Choix du modele : Random Forest')
    para(doc,
        'Nous avons choisi le Random Forest pour trois raisons. Premierement, '
        'il gere nativement les variables categorielles encodees sans necessiter '
        'de normalisation. Deuxiemement, il est robust au bruit et aux valeurs '
        'aberrantes, ce qui est important avec des donnees touristiques reelles. '
        'Troisiemement, le calcul d\'importance des features nous permet d\'expliquer '
        'a la DGTT pourquoi le modele predit ce qu\'il predit, ce qu\'un reseau de '
        'neurones ne permettrait pas facilement.')
    espace(doc)

    titre2(doc, '7.2 Features utilisees')
    tableau(doc,
        ['Feature', 'Type', 'Justification metier'],
        [
            ['mois (1 a 12)', 'Numerique', 'Variable la plus predictive : haute saison novembre-mars (touristes europeens fuyant l\'hiver)'],
            ['destination_enc', 'Categorielle encodee', 'Dakar, Saly et Saint-Louis attirent structurellement plus de visiteurs'],
            ['sentiment_moy', 'Numerique [-1, 1]', 'Une bonne reputation attire davantage de touristes les semaines suivantes'],
            ['note_moy', 'Numerique [2.5, 5.0]', 'Corroboree avec le sentiment pour detecter des incoherences'],
            ['nationalite_enc', 'Categorielle encodee', 'Les Europeens visitent en hiver, les Africains de la sous-region toute l\'annee'],
            ['type_activite_enc', 'Categorielle encodee', 'Le pelerinage a Touba genere des pics tres specifiques'],
        ],
        largeurs=[4, 3.5, 8.5]
    )
    espace(doc)

    titre2(doc, '7.3 Resultats et evaluation')
    tableau(doc,
        ['Metrique', 'Valeur obtenue', 'Interpretation'],
        [
            ['RMSE', '31.23 touristes/jour', 'En moyenne, le modele se trompe de 31 touristes par jour - acceptable pour la planification'],
            ['R2 sur le jeu de test', '0.9005', 'Le modele explique 90% de la variance des flux touristiques'],
            ['R2 en validation croisee (5 plis)', '0.8989', 'L\'ecart avec le R2 test est faible : pas de sur-apprentissage'],
            ['Jeu de donnees', '5000 observations', '80% pour l\'entrainement, 20% pour le test'],
        ],
        largeurs=[5, 4.5, 6.5]
    )
    espace(doc)
    para(doc,
        'Le suivi des experiences est assure par MLflow. Chaque entrainement enregistre '
        'automatiquement les hyperparametres, les metriques et le modele serialise. '
        'Cela permet de comparer les versions et de revenir a un modele anterieur '
        'si une nouvelle version est moins performante.')
    doc.add_page_break()


def airflow(doc):
    titre1(doc, '8. Orchestration : Le DAG Airflow')

    para(doc,
        'Apache Airflow orchestre le cycle de vie du systeme. Le DAG '
        'teranga_sn_monitoring s\'execute chaque lundi a 7h00 du matin '
        'et suit une logique de branchement conditionnel appelee MLOps : '
        'si la qualite des donnees se degrade, le systeme reagit automatiquement.')
    espace(doc)

    titre2(doc, '8.1 Logique de branchement')
    para(doc,
        'La tache check_reputation interroge la vue Hive vue_destinations et compte '
        'le nombre de destinations en statut ROUGE. Cette valeur pilote la suite :')
    espace(doc)
    tableau(doc,
        ['Condition', 'Branche executee', 'Action'],
        [
            ['1 ou 0 destination ROUGE', 'update_dashboard', 'Mise a jour des alertes HBase, pas de reentainement'],
            ['2 destinations ROUGE ou plus', 'retrain_model puis update_dashboard', 'Reentainement du Random Forest via spark-submit, puis mise a jour'],
        ],
        largeurs=[5, 5, 6]
    )
    espace(doc)
    para(doc,
        'Cette logique evite de reentainer le modele a chaque execution inutilement. '
        'Le reentainement n\'est declenche que lorsque la situation sur le terrain '
        'change suffisamment pour que les predictions du modele actuel deviennent '
        'potentiellement obsoletes.')
    espace(doc)

    titre2(doc, '8.2 Taches du DAG')
    tableau(doc,
        ['Tache', 'Type Airflow', 'Description'],
        [
            ['start', 'DummyOperator', 'Point d\'entree du workflow'],
            ['check_reputation', 'BranchPythonOperator', 'Interroge Hive, compte les ROUGE, choisit la branche'],
            ['retrain_model', 'PythonOperator', 'Lance spark-submit sur train_flux_model.py'],
            ['update_dashboard', 'PythonOperator', 'Ecrit les alertes ROUGE/ORANGE dans HBase'],
            ['generer_rapport_semaine', 'PythonOperator', 'Exporte un fichier JSON de synthese des KPIs'],
            ['end', 'DummyOperator', 'Convergence des branches (trigger rule : none_failed)'],
        ],
        largeurs=[4.5, 4.5, 7]
    )
    espace(doc)
    para(doc,
        'L\'interface Airflow est accessible sur http://localhost:8082 '
        '(identifiants : admin / admin). Le DAG apparait avec les tags '
        'teranga-sn, tourisme, ecommerce et mlops.')
    doc.add_page_break()


def dashboard(doc):
    titre1(doc, '9. Tableau de Bord de Decision')

    para(doc,
        'Le tableau de bord est l\'interface finale entre le systeme Big Data et '
        'les decideurs de la DGTT. Il est concu pour etre lisible en quelques '
        'secondes, sans connaissance technique. Il fonctionne en mode demonstration '
        'si Hive n\'est pas disponible.')
    espace(doc)
    tableau(doc,
        ['Panneau', 'Visualisation', 'Lecture rapide pour le decideur'],
        [
            ['Haut gauche', 'Barres horizontales colorees (vert/orange/rouge)', 'D\'un coup d\'oeil : quelles destinations ont des problemes et quelle est leur note'],
            ['Haut droite', 'Barres groupees positif vs negatif', 'Quelle proportion des avis sont negatifs pour chaque destination'],
            ['Bas gauche', 'Barres par categorie e-commerce', 'Quels secteurs generent le plus de chiffre d\'affaires'],
            ['Bas droite', 'Heatmap region x categorie', 'Ou se concentrent les achats par type de produit'],
        ],
        largeurs=[2.5, 5, 8.5]
    )
    espace(doc)
    img = os.path.join(DATA_DIR, 'dashboard_teranga_sn.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        para(doc, 'Figure 1 - Tableau de bord Teranga-SN genere par dashboard_teranga_sn.py', taille=9, italique=True, centre=True)
    doc.add_page_break()


def conclusion(doc):
    titre1(doc, '10. Conclusion et Perspectives')

    para(doc,
        'En trois semaines, l\'equipe Eq.12 a concu et deploye une plateforme Big Data '
        'complete, fonctionnelle et testee, appliquee a un secteur strategique du Senegal. '
        'Chaque composant technique a ete valide individuellement, et le pipeline de bout '
        'en bout a ete teste de l\'ingestion Kafka jusqu\'a l\'ecriture dans HBase.')
    espace(doc)
    para(doc,
        'Ce projet nous a appris plusieurs lecons importantes. D\'abord, que la '
        'complexite en Big Data vient rarement des algorithmes eux-memes, mais de '
        'la coordination entre les services. Faire communiquer Kafka, Spark et HBase '
        'dans un reseau Docker demande de comprendre les protocols de chaque outil. '
        'Ensuite, que la qualite des donnees doit etre traitee des le debut du pipeline, '
        'pas a la fin. Les tests Pandera nous ont permis de detecter des erreurs '
        'd\'anonymisation qui auraient pu passer inapercues.')
    espace(doc)

    titre2(doc, '10.1 Bilan technique')
    tableau(doc,
        ['Livrable', 'Etat'],
        [
            ['8 services Docker integres et valides', 'Termine'],
            ['Pipeline Kafka : 2 topics, producteur Python + NiFi', 'Termine'],
            ['Spark Streaming : NLP + SHA-256 + watermark 1 jour', 'Termine'],
            ['HBase : 3 tables dans namespace teranga, lecture/ecriture validees', 'Termine'],
            ['Hive : base teranga_sn, 2 tables ORC/SNAPPY, 2 vues analytiques', 'Termine'],
            ['Pandera : 10 tests de validation de schema (10/10)', 'Termine'],
            ['Random Forest : RMSE=31.23, R2=0.90, suivi MLflow', 'Termine'],
            ['Airflow DAG : monitoring hebdomadaire avec branchement MLOps', 'Termine'],
            ['Dashboard : 4 panneaux de visualisation generes en PNG', 'Termine'],
            ['Rapport final et documentation binome (TINE_ONBOARDING.md)', 'Termine'],
        ],
        largeurs=[12, 4]
    )
    espace(doc)

    titre2(doc, '10.2 Axes d\'amelioration identifies')
    tableau(doc,
        ['Axe', 'Amelioration envisagee'],
        [
            ['Lexique Wolof', 'Enrichissement avec des locuteurs natifs pour couvrir les expressions idiomatiques et les variations dialectales'],
            ['Modele NLP', 'Remplacer le lexique par un modele transformeur (CamemBERT ou XLM-RoBERTa) pour le francais et l\'anglais'],
            ['HBase TTL', 'Configurer l\'expiration automatique des alertes apres 7 jours pour ne pas saturer la memoire'],
            ['NiFi REST', 'Connecter NiFi directement a l\'API REST TripAdvisor pour remplacer les donnees simulees par des donnees reelles'],
            ['Scalabilite', 'Migrer vers Kubernetes pour gerer des pics de charge lors des grands evenements touristiques'],
        ],
        largeurs=[4, 12]
    )
    espace(doc)
    para(doc,
        'Le code source complet est disponible sur le depot GitHub : '
        'https://github.com/pa-malick/uadb-m2-teranga-sn',
        italique=True
    )


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    page_titre(doc)
    resume(doc)
    introduction(doc)
    architecture(doc)
    kafka_nifi(doc)
    spark(doc)
    stockage(doc)
    validation(doc)
    machine_learning(doc)
    airflow(doc)
    dashboard(doc)
    conclusion(doc)

    doc.save(OUT_PATH)
    print(f'Rapport genere : {OUT_PATH}')
    print(f'Taille : {os.path.getsize(OUT_PATH) // 1024} Ko')


if __name__ == '__main__':
    main()
